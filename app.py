import hashlib
import hmac
import io
import secrets as py_secrets
from datetime import datetime, timedelta, timezone
from typing import Optional

import pandas as pd
import requests
import streamlit as st

st.set_page_config(page_title="Job Matrix Manager", layout="wide")

st.markdown(
    """
<style>
.block-container {padding-top: 1rem; padding-bottom: 1.5rem; max-width: 1450px;}
html, body, [class*="css"] {font-family: Arial, Helvetica, sans-serif;}
.handout-box {white-space: pre-wrap; border: 1px solid #d9dde3; padding: 1rem 1.1rem; background: #ffffff; border-radius: 0.35rem; line-height: 1.8;}
.small-note {font-size: 0.86rem; color: #4b5563;}
div[data-testid="stMetric"] {border: 1px solid #d9dde3; padding: 0.8rem; border-radius: 0.35rem; background: #ffffff;}
</style>
""",
    unsafe_allow_html=True,
)

st.title("Job Matrix Manager")
st.caption("部門・グレード別の職務定義を管理し、CSV / Excel / Word / PDF 出力できます。専用ログイン、5回失敗ロック、セッション期限対応。")

# =========================
# SECRETS / REST SETTINGS (STRICT JM ONLY)
# =========================
required_secret_keys = [
    "JM_SUPABASE_URL",
    "JM_SUPABASE_SERVICE_ROLE_KEY",
]

missing_keys = [k for k in required_secret_keys if not str(st.secrets.get(k, "")).strip()]
if missing_keys:
    st.error(
        "Missing required Streamlit secrets: "
        + ", ".join(missing_keys)
        + ". This app only accepts JM_* secrets."
    )
    st.stop()

SUPABASE_URL = str(st.secrets["JM_SUPABASE_URL"]).rstrip("/")
SUPABASE_KEY = str(st.secrets["JM_SUPABASE_SERVICE_ROLE_KEY"]).strip()

APP_NAME = st.secrets.get("JM_APP_NAME", "Job Matrix Manager")
USERS_TABLE = st.secrets.get("JM_USERS_TABLE", "job_matrix_users")
SNAPSHOTS_TABLE = st.secrets.get("JM_SNAPSHOTS_TABLE", "job_matrix_snapshots")
LOGIN_ID_LABEL = st.secrets.get("JM_LOGIN_ID_LABEL", "User ID")
LOCK_MAX_ATTEMPTS = int(st.secrets.get("JM_LOCK_MAX_ATTEMPTS", 5))
LOCK_MINUTES = int(st.secrets.get("JM_LOCK_MINUTES", 30))
SESSION_TIMEOUT_HOURS = int(st.secrets.get("JM_SESSION_TIMEOUT_HOURS", 12))


def now_utc() -> datetime:
    return datetime.now(timezone.utc)


def supabase_headers(prefer: Optional[str] = None) -> dict:
    headers = {
        "apikey": SUPABASE_KEY,
        "Authorization": f"Bearer {SUPABASE_KEY}",
        "Content-Type": "application/json",
    }
    if prefer:
        headers["Prefer"] = prefer
    return headers


def rest_get(path: str, params: Optional[dict] = None):
    response = requests.get(
        f"{SUPABASE_URL}/rest/v1/{path}",
        headers=supabase_headers(),
        params=params,
        timeout=30,
    )
    response.raise_for_status()
    return response.json()


def rest_post(path: str, payload):
    response = requests.post(
        f"{SUPABASE_URL}/rest/v1/{path}",
        headers=supabase_headers(prefer="return=representation"),
        json=payload,
        timeout=30,
    )
    response.raise_for_status()
    return response.json()


def rest_patch(path: str, payload: dict, params: Optional[dict] = None):
    response = requests.patch(
        f"{SUPABASE_URL}/rest/v1/{path}",
        headers=supabase_headers(prefer="return=representation"),
        params=params,
        json=payload,
        timeout=30,
    )
    response.raise_for_status()
    return response.json()


def rest_delete(path: str, params: Optional[dict] = None):
    response = requests.delete(
        f"{SUPABASE_URL}/rest/v1/{path}",
        headers=supabase_headers(prefer="return=representation"),
        params=params,
        timeout=30,
    )
    response.raise_for_status()
    return response.json() if response.text else []


# =========================
# AUTH HELPERS
# =========================
def hash_password(password: str) -> str:
    return hashlib.sha256(password.encode("utf-8")).hexdigest()


def verify_password(password: str, password_hash: str) -> bool:
    return hmac.compare_digest(hash_password(password), str(password_hash or ""))


def normalize_username(username: str) -> str:
    return str(username).strip()


def get_user_by_username(username: str) -> Optional[dict]:
    users = rest_get(
        USERS_TABLE,
        params={
            "select": "*",
            "username": f"eq.{normalize_username(username)}",
            "limit": 1,
        },
    )
    return users[0] if users else None


def parse_locked_until(value) -> Optional[datetime]:
    if not value:
        return None
    try:
        dt = datetime.fromisoformat(str(value).replace("Z", "+00:00"))
        if dt.tzinfo is None:
            dt = dt.replace(tzinfo=timezone.utc)
        return dt.astimezone(timezone.utc)
    except Exception:
        return None


def is_user_locked(user: dict) -> tuple[bool, Optional[datetime]]:
    locked_until = parse_locked_until(user.get("locked_until"))
    if locked_until and locked_until > now_utc():
        return True, locked_until
    return False, locked_until


def update_user_row(username: str, payload: dict):
    return rest_patch(USERS_TABLE, payload, params={"username": f"eq.{normalize_username(username)}"})


def register_failed_attempt(username: str, current_failed_attempts: int):
    next_attempts = int(current_failed_attempts or 0) + 1
    payload = {"failed_attempts": next_attempts, "updated_at": now_utc().isoformat()}
    if next_attempts >= LOCK_MAX_ATTEMPTS:
        payload["locked_until"] = (now_utc() + timedelta(minutes=LOCK_MINUTES)).isoformat()
    update_user_row(username, payload)


def clear_failed_attempts(username: str):
    update_user_row(
        username,
        {
            "failed_attempts": 0,
            "locked_until": None,
            "updated_at": now_utc().isoformat(),
        },
    )

def session_expired() -> bool:
    expires_at = st.session_state.get("auth_expires_at")
    if not expires_at:
        return True
    parsed = parse_locked_until(expires_at)
    if parsed is None:
        return True
    return parsed <= now_utc()


def clear_auth_session():
    for key in [
        "auth_user",
        "auth_expires_at",
        "df",
        "grade_master",
        "last_change_log",
        "last_grade_change_log",
    ]:
        if key in st.session_state:
            del st.session_state[key]


def set_auth_session(user: dict):
    st.session_state.auth_user = {
        "username": user["username"],
        "display_name": user.get("display_name") or user["username"],
        "role": user.get("role", "viewer"),
    }
    st.session_state.auth_expires_at = (now_utc() + timedelta(hours=SESSION_TIMEOUT_HOURS)).isoformat()


def require_login() -> dict:
    st.subheader("ログイン")
    st.caption("管理人は編集可能、ユーザーは閲覧・ダウンロードのみです。パスワード5回失敗で30分ロックされます。")

    with st.form("login_form"):
        username = st.text_input(LOGIN_ID_LABEL)
        password = st.text_input("パスワード", type="password")
        submitted = st.form_submit_button("ログイン", type="primary")

    if not submitted:
        st.stop()

    user = get_user_by_username(username)
    if not user:
        st.error("ユーザー名またはパスワードが違います。")
        st.stop()

    if not bool(user.get("is_active", True)):
        st.error("このユーザーは無効化されています。")
        st.stop()

    locked, locked_until = is_user_locked(user)
    if locked:
        st.error(f"このアカウントはロック中です。解除予定: {locked_until.astimezone().strftime('%Y-%m-%d %H:%M:%S')}")
        st.stop()

    if not verify_password(password, user.get("password_hash", "")):
        register_failed_attempt(user["username"], int(user.get("failed_attempts") or 0))
        remain = max(0, LOCK_MAX_ATTEMPTS - (int(user.get("failed_attempts") or 0) + 1))
        if remain > 0:
            st.error(f"ユーザー名またはパスワードが違います。残り {remain} 回でロックされます。")
        else:
            st.error(f"パスワード誤入力が {LOCK_MAX_ATTEMPTS} 回に達したため、30分ロックしました。")
        st.stop()

    clear_failed_attempts(user["username"])
    set_auth_session(user)
    st.rerun()


# =========================
# MASTER DATA
# =========================
GRADES = {
    "G6": "ジュニアスタッフ",
    "G5B": "シニアスタッフ",
    "G5A": "上位シニアスタッフ",
    "G4": "スーパーバイザー",
    "G3": "課長",
    "G2": "次長",
}
GRADE_ORDER = ["G6", "G5B", "G5A", "G4", "G3", "G2"]
DEPARTMENTS = [
    "人事課", "総務課", "経理課", "財務課", "情報管理課", "エリア運営課",
    "設備・運搬課", "ナーサリー課", "マーケティング課", "プロセス管理課",
    "インサイドコミュニケーション課", "アウトサイドコミュニケーション課",
]
UPDATE_COLUMNS = [
    "grade_name", "role_summary", "kpi", "reports_to", "direct_reports",
    "skills", "kra", "accountabilities",
]
DEFAULT_GRADE_OVERVIEW = {
    "G6": "上位者の指示に基づき担当業務を正確かつ着実に遂行するグレードである。",
    "G5B": "一定の自律性をもって担当業務を安定運用し、周囲と連携して成果に貢献するグレードである。",
    "G5A": "担当領域を主体的に遂行し、業務品質の向上や改善提案にも関与するグレードである。",
    "G4": "課単位の目標達成に責任を持ち、業務と人材の管理を通じて成果最大化を担うグレードである。",
    "G3": "事業部単位の計画立案・組織運営・人材マネジメントを統括するグレードである。",
    "G2": "拠点全体の戦略・資源配分・組織運営を統括し、全体最適を牽引するグレードである。",
}
DEPARTMENT_DESCRIPTIONS = {
    "人事課": "採用・労務・人材配置・教育支援を担う。",
    "総務課": "拠点運営を支える庶務・文書・社内環境整備を担う。",
    "経理課": "日常経理、締め処理、会計記録を担う。",
    "財務課": "資金繰り、予算管理、金融機関対応を担う。",
    "情報管理課": "システム運用、データ整備、情報活用を担う。",
    "エリア運営課": "農地・現場の運営管理、作業調整、巡回確認を担う。",
    "設備・運搬課": "設備保守、資材運搬、車両・物流運用を担う。",
    "ナーサリー課": "苗木育成、品質管理、植栽前の育苗管理を担う。",
    "マーケティング課": "市場情報整理、外部向け資料、事業訴求支援を担う。",
    "プロセス管理課": "業務手順、工程、品質、安全の標準化を担う。",
    "インサイドコミュニケーション課": "社内周知、社内文化、組織浸透施策を担う。",
    "アウトサイドコミュニケーション課": "対外発信、渉外、地域・関係者連携を担う。",
}


def default_reports_to(grade: str) -> str:
    mapping = {
        "G6": "担当課の上位者またはスーパーバイザー",
        "G5B": "担当課の上位者またはスーパーバイザー",
        "G5A": "担当課のスーパーバイザーまたは課長",
        "G4": "課長 / 部門責任者",
        "G3": "次長 / 部門統括責任者",
        "G2": "拠点長 / 経営層",
    }
    return mapping[grade]


def default_direct_reports(grade: str) -> str:
    mapping = {
        "G6": "なし",
        "G5B": "なし、または後輩スタッフへの実務支援",
        "G5A": "小規模な実務リード、または後輩スタッフ支援",
        "G4": "担当チームメンバー、現場スタッフ",
        "G3": "複数チーム、各課の責任者、スーパーバイザー",
        "G2": "複数部門責任者、課長クラス",
    }
    return mapping[grade]


def default_skills(grade: str) -> str:
    if grade in ["G6", "G5B", "G5A"]:
        return "Excel / Word / メール、データ入力、報告連携、基本的な業務理解"
    if grade == "G4":
        return "Excel集計、進捗管理、レポーティング、メンバー指導、部門調整"
    if grade == "G3":
        return "計画立案、予実管理、KPI管理、部門横断調整、人材マネジメント"
    return "戦略立案、組織運営、数値管理、意思決定、経営報告"


def default_kra(grade: str) -> str:
    if grade in ["G6", "G5B"]:
        return "担当業務の安定運用、処理品質の維持、期限遵守"
    if grade == "G5A":
        return "担当領域の安定運用、品質向上、業務改善"
    if grade == "G4":
        return "チーム成果、人員配置、業務改善、コスト最適化"
    if grade == "G3":
        return "部門目標達成、予算管理、運営改善、人材育成"
    return "拠点目標達成、全体最適、収益改善、組織強化"


def default_accountabilities(grade: str) -> str:
    if grade in ["G6", "G5B", "G5A"]:
        return "日次業務処理、定型レポート作成、記録更新、関連部門連携"
    if grade == "G4":
        return "進捗管理、会議運営、月次報告、メンバー指導、改善提案"
    if grade == "G3":
        return "部門報告、予実レビュー、方針展開、課題管理、承認対応"
    return "重要課題レビュー、経営報告、部門統括、資源配分、意思決定"


def make_default_role_summary(department: str, grade: str, grade_name: str) -> str:
    return f"{grade_name}として{department}の業務を担う。担当領域の業務を正確かつ継続的に遂行し、関係者と連携しながら成果達成に貢献する。必要に応じて改善や調整にも関わり、組織運営を支える役割を果たす。"


def make_default_kpi(grade: str) -> str:
    return {
        "G6": "業務達成率、処理件数、正確性、期限遵守率",
        "G5B": "担当業務達成率、品質、改善件数、納期遵守率",
        "G5A": "成果達成率、改善提案数、関連部署連携品質",
        "G4": "チーム達成率、進捗遵守率、品質指標、人材育成状況",
        "G3": "課目標達成率、予算遵守率、改善成果、組織運営状況",
        "G2": "部門目標達成率、戦略進捗、横断課題解決、収益・効率改善",
    }[grade]


def get_default_grade_master() -> pd.DataFrame:
    return pd.DataFrame([
        {"grade": grade, "grade_name": GRADES[grade], "grade_overview": DEFAULT_GRADE_OVERVIEW[grade]}
        for grade in GRADE_ORDER
    ])


def generate_default() -> pd.DataFrame:
    rows = []
    for dept in DEPARTMENTS:
        for grade in GRADE_ORDER:
            rows.append({
                "department": dept,
                "grade": grade,
                "grade_name": GRADES[grade],
                "role_summary": make_default_role_summary(dept, grade, GRADES[grade]),
                "kpi": make_default_kpi(grade),
                "reports_to": default_reports_to(grade),
                "direct_reports": default_direct_reports(grade),
                "skills": default_skills(grade),
                "kra": default_kra(grade),
                "accountabilities": default_accountabilities(grade),
            })
    return pd.DataFrame(rows)


def ensure_main_columns(df: pd.DataFrame) -> pd.DataFrame:
    out = df.copy()
    defaults = {
        "department": "",
        "grade": "",
        "grade_name": "",
        "role_summary": "",
        "kpi": "",
        "reports_to": "",
        "direct_reports": "",
        "skills": "",
        "kra": "",
        "accountabilities": "",
    }
    for col, default in defaults.items():
        if col not in out.columns:
            out[col] = default
    return out[list(defaults.keys())]


def ensure_grade_master(df: Optional[pd.DataFrame]) -> pd.DataFrame:
    base = get_default_grade_master()
    if df is None or df.empty:
        return base
    out = df.copy()
    for col in ["grade", "grade_name", "grade_overview"]:
        if col not in out.columns:
            out[col] = ""
    out = out[["grade", "grade_name", "grade_overview"]].copy()
    out["grade"] = out["grade"].astype(str).str.strip()
    out = out.drop_duplicates(subset=["grade"], keep="last")
    merged = base[["grade"]].merge(out, on="grade", how="left")
    merged["grade_name"] = merged["grade_name"].replace("", pd.NA).fillna(merged["grade"].map(GRADES))
    merged["grade_overview"] = merged["grade_overview"].replace("", pd.NA).fillna(merged["grade"].map(DEFAULT_GRADE_OVERVIEW))
    merged["grade_sort"] = merged["grade"].map({g: i for i, g in enumerate(GRADE_ORDER)})
    return merged.sort_values("grade_sort").drop(columns=["grade_sort"])


def attach_grade_overview(main_df: pd.DataFrame, grade_master: pd.DataFrame) -> pd.DataFrame:
    df = ensure_main_columns(main_df)
    gm = ensure_grade_master(grade_master)
    merged = df.merge(gm[["grade", "grade_overview"]], on="grade", how="left")
    merged["grade_overview"] = merged["grade_overview"].fillna(merged["grade"].map(DEFAULT_GRADE_OVERVIEW))
    return merged


def normalize_text(value):
    if pd.isna(value):
        return None
    text = str(value).strip()
    return text if text != "" else None


def validate_import_keys(base_df: pd.DataFrame, upload_df: pd.DataFrame) -> pd.DataFrame:
    base_keys = set(base_df["department"].astype(str).str.strip() + "||" + base_df["grade"].astype(str).str.strip())
    up = upload_df.copy()
    up["_merge_key"] = up["department"].astype(str).str.strip() + "||" + up["grade"].astype(str).str.strip()
    invalid = up[~up["_merge_key"].isin(base_keys)].copy()
    return invalid.drop(columns=["_merge_key"])


def merge_partial_update(base_df: pd.DataFrame, upload_df: pd.DataFrame):
    required_keys = ["department", "grade"]
    for col in required_keys:
        if col not in upload_df.columns:
            raise ValueError(f"CSVに必須列 '{col}' がありません。")

    updated_df = ensure_main_columns(base_df.copy())
    import_df = upload_df.copy()

    for col in required_keys:
        updated_df[col] = updated_df[col].astype(str).str.strip()
        import_df[col] = import_df[col].astype(str).str.strip()

    import_df = import_df.drop_duplicates(subset=required_keys, keep="last")
    updated_df["_merge_key"] = updated_df["department"] + "||" + updated_df["grade"]
    import_df["_merge_key"] = import_df["department"] + "||" + import_df["grade"]
    import_map = import_df.set_index("_merge_key").to_dict(orient="index")

    changes = []
    for idx, row in updated_df.iterrows():
        key = row["_merge_key"]
        if key not in import_map:
            continue
        source = import_map[key]
        for col in UPDATE_COLUMNS:
            if col not in import_df.columns:
                continue
            old_val = normalize_text(row.get(col))
            new_val = normalize_text(source.get(col))
            if new_val is not None and new_val != old_val:
                updated_df.at[idx, col] = new_val
                changes.append({
                    "department": row["department"],
                    "grade": row["grade"],
                    "column": col,
                    "before": old_val,
                    "after": new_val,
                })

    updated_df = updated_df.drop(columns=["_merge_key"])
    return updated_df, pd.DataFrame(changes)


def merge_grade_master_update(base_df: pd.DataFrame, upload_df: pd.DataFrame):
    if "grade" not in upload_df.columns:
        raise ValueError("CSVに必須列 'grade' がありません。")

    updated_df = ensure_grade_master(base_df.copy())
    import_df = upload_df.copy()
    for col in ["grade", "grade_name", "grade_overview"]:
        if col not in import_df.columns:
            import_df[col] = None
    updated_df["grade"] = updated_df["grade"].astype(str).str.strip()
    import_df["grade"] = import_df["grade"].astype(str).str.strip()
    import_df = import_df.drop_duplicates(subset=["grade"], keep="last")
    import_map = import_df.set_index("grade").to_dict(orient="index")

    changes = []
    for idx, row in updated_df.iterrows():
        grade = row["grade"]
        if grade not in import_map:
            continue
        source = import_map[grade]
        for col in ["grade_name", "grade_overview"]:
            old_val = normalize_text(row.get(col))
            new_val = normalize_text(source.get(col))
            if new_val is not None and new_val != old_val:
                updated_df.at[idx, col] = new_val
                changes.append({"grade": grade, "column": col, "before": old_val, "after": new_val})
    return updated_df, pd.DataFrame(changes)


def update_master_from_editor(master_df: pd.DataFrame, edited_df: pd.DataFrame) -> pd.DataFrame:
    current = ensure_main_columns(master_df.copy())
    edited = ensure_main_columns(edited_df.copy())
    edited_map = edited.set_index(["department", "grade"]).to_dict(orient="index")
    for idx, row in current.iterrows():
        key = (row["department"], row["grade"])
        if key in edited_map:
            for col in UPDATE_COLUMNS:
                current.at[idx, col] = edited_map[key].get(col, row[col])
    return current


def department_summary(df: pd.DataFrame) -> pd.DataFrame:
    rows = []
    df = ensure_main_columns(df.copy())
    for dept in DEPARTMENTS:
        sub = df[df["department"] == dept]
        rows.append({
            "department": dept,
            "rows": len(sub),
            "filled_role_summary": int(sub["role_summary"].astype(str).str.strip().ne("").sum()),
            "filled_kpi": int(sub["kpi"].astype(str).str.strip().ne("").sum()),
            "filled_kra": int(sub["kra"].astype(str).str.strip().ne("").sum()),
        })
    return pd.DataFrame(rows)


def split_kpi_items(kpi_text: str) -> list[str]:
    if pd.isna(kpi_text) or str(kpi_text).strip() == "":
        return []
    text = str(kpi_text).replace("\n", "、")
    parts = [p.strip(" ・-\t") for p in text.replace(",", "、").split("、")]
    return [p for p in parts if p]


def make_handout_dict(row: pd.Series) -> dict:
    department = str(row.get("department", "")).strip()
    grade = str(row.get("grade", "")).strip()
    grade_name = str(row.get("grade_name", "")).strip()
    role_summary = str(row.get("role_summary", "")).strip()
    grade_overview = str(row.get("grade_overview", "")).strip()
    kpi_items = split_kpi_items(str(row.get("kpi", "")))
    return {
        "title": "配属・職務定義書",
        "department": department,
        "position": f"{grade}（{grade_name}）",
        "job_overview": role_summary if role_summary else "担当領域の業務を遂行します。",
        "grade_definition": f"本職位は、{grade_overview}" if grade_overview and not grade_overview.startswith("本職位は") else grade_overview,
        "kpi_items": kpi_items,
        "reports_to": str(row.get("reports_to", "")).strip(),
        "direct_reports": str(row.get("direct_reports", "")).strip(),
        "skills": str(row.get("skills", "")).strip(),
        "kra": str(row.get("kra", "")).strip(),
        "accountabilities": str(row.get("accountabilities", "")).strip(),
        "note": "本定義は、業務内容・組織状況に応じて見直される場合があります。",
    }


def make_handout_text(row: pd.Series) -> str:
    data = make_handout_dict(row)
    kpi_text = "\n".join([f"・{item}" for item in data["kpi_items"]]) if data["kpi_items"] else "・別途設定"
    return (
        f"【{data['title']}】\n\n"
        f"■ 所属\n{data['department']}\n\n"
        f"■ 職位\n{data['position']}\n\n"
        f"■ Reports to（直属の上司）\n{data['reports_to'] or '別途設定'}\n\n"
        f"■ Direct Reports（直属の部下）\n{data['direct_reports'] or '別途設定'}\n\n"
        f"■ 職務概要\n{data['job_overview']}\n\n"
        f"■ Skills\n{data['skills'] or '別途設定'}\n\n"
        f"■ Key Result Areas - KRA（主要な責任）\n{data['kra'] or '別途設定'}\n\n"
        f"■ Key Operational Accountabilities（重要なタスク）\n{data['accountabilities'] or '別途設定'}\n\n"
        f"■ グレード定義\n{data['grade_definition']}\n\n"
        f"■ 評価指標（KPI）\n{kpi_text}\n\n"
        f"■ 備考\n{data['note']}"
    )


def dataframe_to_csv_bytes(df: pd.DataFrame) -> bytes:
    return ensure_main_columns(df.copy()).to_csv(index=False).encode("utf-8-sig")


def grade_master_to_csv_bytes(df: pd.DataFrame) -> bytes:
    return ensure_grade_master(df.copy()).to_csv(index=False).encode("utf-8-sig")


def minimal_template_csv_bytes() -> bytes:
    return pd.DataFrame(columns=[
        "department", "grade", "grade_name", "role_summary", "kpi", "reports_to",
        "direct_reports", "skills", "kra", "accountabilities",
    ]).to_csv(index=False).encode("utf-8-sig")


def minimal_grade_template_csv_bytes() -> bytes:
    return pd.DataFrame(columns=["grade", "grade_name", "grade_overview"]).to_csv(index=False).encode("utf-8-sig")


def dataframe_to_excel_bytes(df: pd.DataFrame, grade_master_df: pd.DataFrame) -> bytes:
    output = io.BytesIO()
    export_df = attach_grade_overview(df.copy(), grade_master_df.copy())
    export_df["handout_preview"] = export_df.apply(make_handout_text, axis=1)
    with pd.ExcelWriter(output, engine="openpyxl") as writer:
        export_df.to_excel(writer, index=False, sheet_name=SNAPSHOTS_TABLE)
        ensure_grade_master(grade_master_df.copy()).to_excel(writer, index=False, sheet_name="grade_master")
    output.seek(0)
    return output.getvalue()


def dataframe_for_handout_export(df: pd.DataFrame) -> pd.DataFrame:
    out = df.copy()
    out["grade_sort"] = out["grade"].map({g: i for i, g in enumerate(GRADE_ORDER)})
    return out.sort_values(["department", "grade_sort"]).drop(columns=["grade_sort"])


def create_handout_excel_bytes(export_df: pd.DataFrame) -> bytes:
    from openpyxl import Workbook
    from openpyxl.styles import Alignment, Border, Font, PatternFill, Side

    wb = Workbook()
    ws0 = wb.active
    wb.remove(ws0)

    thin = Side(style="thin", color="C7D2E0")
    border = Border(left=thin, right=thin, top=thin, bottom=thin)
    fill = PatternFill("solid", fgColor="EAF1FB")
    title_font = Font(name="Arial", size=14, bold=True)
    body_font = Font(name="Arial", size=10.5)
    label_font = Font(name="Arial", size=10.5, bold=True)

    for i, record in enumerate(export_df.to_dict(orient="records"), start=1):
        row = pd.Series(record)
        data = make_handout_dict(row)
        ws = wb.create_sheet(title=f"{str(data['department'])[:20]}_{row['grade']}")
        ws.column_dimensions["A"].width = 25
        ws.column_dimensions["B"].width = 70
        ws.merge_cells("A1:B1")
        ws["A1"] = data["title"]
        ws["A1"].font = title_font
        ws["A1"].alignment = Alignment(horizontal="center", vertical="center")
        rows = [
            ("所属", data["department"]),
            ("職位", data["position"]),
            ("直属の上司", data["reports_to"] or "別途設定"),
            ("直属の部下", data["direct_reports"] or "別途設定"),
            ("職務概要", data["job_overview"]),
            ("必要スキル", data["skills"] or "別途設定"),
            ("KRA（主要な責任）", data["kra"] or "別途設定"),
            ("重要タスク", data["accountabilities"] or "別途設定"),
            ("グレード定義", data["grade_definition"]),
            ("評価指標（KPI）", "\n".join([f"・{x}" for x in data["kpi_items"]]) if data["kpi_items"] else "・別途設定"),
            ("備考", data["note"]),
        ]
        start = 3
        for offset, (label, value) in enumerate(rows):
            r = start + offset
            ws[f"A{r}"] = label
            ws[f"B{r}"] = value
            ws[f"A{r}"].fill = fill
            ws[f"A{r}"].font = label_font
            ws[f"B{r}"].font = body_font
            ws[f"A{r}"].border = border
            ws[f"B{r}"].border = border
            ws[f"A{r}"].alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
            ws[f"B{r}"].alignment = Alignment(horizontal="left", vertical="top", wrap_text=True)
        ws[f"B{start + len(rows) + 1}"] = f"作成日: {datetime.now().strftime('%Y-%m-%d')}"

    buffer = io.BytesIO()
    wb.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()


def create_handout_docx_bytes(export_df: pd.DataFrame) -> bytes:
    from docx import Document
    from docx.enum.section import WD_SECTION
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt

    doc = Document()
    for idx, record in enumerate(export_df.to_dict(orient="records")):
        row = pd.Series(record)
        data = make_handout_dict(row)
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(data["title"])
        run.bold = True
        run.font.size = Pt(15)
        table = doc.add_table(rows=0, cols=2)
        table.style = "Table Grid"
        rows = [
            ("所属", data["department"]),
            ("職位", data["position"]),
            ("直属の上司", data["reports_to"] or "別途設定"),
            ("直属の部下", data["direct_reports"] or "別途設定"),
            ("職務概要", data["job_overview"]),
            ("必要スキル", data["skills"] or "別途設定"),
            ("KRA（主要な責任）", data["kra"] or "別途設定"),
            ("重要タスク", data["accountabilities"] or "別途設定"),
            ("グレード定義", data["grade_definition"]),
            ("評価指標（KPI）", "\n".join([f"・{x}" for x in data["kpi_items"]]) if data["kpi_items"] else "・別途設定"),
            ("備考", data["note"]),
        ]
        for label, value in rows:
            cells = table.add_row().cells
            cells[0].text = label
            cells[1].text = value
        if idx != len(export_df) - 1:
            doc.add_section(WD_SECTION.NEW_PAGE)
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()


def create_handout_pdf_bytes(export_df: pd.DataFrame) -> bytes:
    from reportlab.lib import colors
    from reportlab.lib.enums import TA_CENTER, TA_LEFT
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
    from reportlab.lib.units import mm
    from reportlab.pdfbase.cidfonts import UnicodeCIDFont
    from reportlab.pdfbase.pdfmetrics import registerFont
    from reportlab.platypus import PageBreak, Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle

    registerFont(UnicodeCIDFont("HeiseiKakuGo-W5"))
    registerFont(UnicodeCIDFont("HeiseiMin-W3"))
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=A4, leftMargin=16*mm, rightMargin=16*mm, topMargin=15*mm, bottomMargin=15*mm)
    styles = getSampleStyleSheet()
    title_style = ParagraphStyle("TitleJP", parent=styles["Title"], fontName="HeiseiKakuGo-W5", fontSize=14, leading=18, alignment=TA_CENTER)
    body_style = ParagraphStyle("BodyJP", parent=styles["BodyText"], fontName="HeiseiMin-W3", fontSize=10, leading=15, alignment=TA_LEFT)
    label_style = ParagraphStyle("LabelJP", parent=body_style, fontName="HeiseiKakuGo-W5")
    story = []
    records = export_df.to_dict(orient="records")
    for idx, record in enumerate(records):
        row = pd.Series(record)
        data = make_handout_dict(row)
        story.append(Paragraph(data["title"], title_style))
        story.append(Spacer(1, 3*mm))
        rows = [
            [Paragraph("所属", label_style), Paragraph(data["department"], body_style)],
            [Paragraph("職位", label_style), Paragraph(data["position"], body_style)],
            [Paragraph("直属の上司", label_style), Paragraph(data["reports_to"] or "別途設定", body_style)],
            [Paragraph("直属の部下", label_style), Paragraph(data["direct_reports"] or "別途設定", body_style)],
            [Paragraph("職務概要", label_style), Paragraph(data["job_overview"].replace("\n", "<br/>"), body_style)],
            [Paragraph("必要スキル", label_style), Paragraph((data["skills"] or "別途設定").replace("\n", "<br/>"), body_style)],
            [Paragraph("KRA（主要な責任）", label_style), Paragraph((data["kra"] or "別途設定").replace("\n", "<br/>"), body_style)],
            [Paragraph("重要タスク", label_style), Paragraph((data["accountabilities"] or "別途設定").replace("\n", "<br/>"), body_style)],
            [Paragraph("グレード定義", label_style), Paragraph(data["grade_definition"].replace("\n", "<br/>"), body_style)],
            [Paragraph("評価指標（KPI）", label_style), Paragraph("<br/>".join([f"・{x}" for x in data["kpi_items"]]) if data["kpi_items"] else "・別途設定", body_style)],
            [Paragraph("備考", label_style), Paragraph(data["note"], body_style)],
        ]
        table = Table(rows, colWidths=[40*mm, 134*mm])
        table.setStyle(TableStyle([
            ("BACKGROUND", (0, 0), (0, -1), colors.HexColor("#EAF1FB")),
            ("BOX", (0, 0), (-1, -1), 0.75, colors.HexColor("#C7D2E0")),
            ("INNERGRID", (0, 0), (-1, -1), 0.5, colors.HexColor("#C7D2E0")),
            ("VALIGN", (0, 0), (-1, -1), "TOP"),
            ("LEFTPADDING", (0, 0), (-1, -1), 8),
            ("RIGHTPADDING", (0, 0), (-1, -1), 8),
            ("TOPPADDING", (0, 0), (-1, -1), 7),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 7),
        ]))
        story.append(table)
        if idx != len(records) - 1:
            story.append(PageBreak())
    doc.build(story)
    buffer.seek(0)
    return buffer.getvalue()


def save_snapshot(main_df: pd.DataFrame, grade_master_df: pd.DataFrame, memo: str) -> None:
    payload = {
        "data": ensure_main_columns(main_df).to_dict(orient="records"),
        "grade_master": ensure_grade_master(grade_master_df).to_dict(orient="records"),
        "memo": memo,
        "created_at": now_utc().isoformat(),
    }
    rest_post(SNAPSHOTS_TABLE, payload)


def load_history() -> list:
    return rest_get(SNAPSHOTS_TABLE, params={"select": "*", "order": "created_at.desc"})


def load_latest_snapshot():
    history = load_history()
    if not history:
        return None, None
    latest = history[0]
    return ensure_main_columns(pd.DataFrame(latest.get("data", []))), ensure_grade_master(pd.DataFrame(latest.get("grade_master", [])))


# =========================
# CONNECTION CHECK
# =========================
if "connection_ok" not in st.session_state:
    try:
        _ = load_history()
        _ = rest_get(USERS_TABLE, params={"select": "username", "limit": 1})
        st.session_state.connection_ok = True
    except Exception as e:
        st.session_state.connection_ok = False
        st.session_state.connection_error = repr(e)

if not st.session_state.connection_ok:
    st.error(f"Supabase REST 接続失敗: {st.session_state.get('connection_error', 'Unknown error')}")
    st.info("JM_* 系の Secrets、専用テーブル、service role key を確認してください。")
    st.stop()

if "auth_user" in st.session_state and session_expired():
    clear_auth_session()
    st.warning("セッションの有効期限が切れたため、再ログインしてください。")

if "auth_user" not in st.session_state:
    require_login()
else:
    st.session_state.auth_expires_at = (now_utc() + timedelta(hours=SESSION_TIMEOUT_HOURS)).isoformat()

auth_user = st.session_state["auth_user"]
is_admin = auth_user.get("role") == "admin"

with st.sidebar:
    st.markdown(f"### {APP_NAME}")
    st.write(f"ログイン中: **{auth_user.get('display_name')}**")
    st.write(f"権限: **{auth_user.get('role')}**")
    st.caption(f"セッション有効時間: {SESSION_TIMEOUT_HOURS} 時間")
    if st.button("ログアウト"):
        clear_auth_session()
        st.rerun()

# =========================
# SESSION DATA
# =========================
if "df" not in st.session_state or "grade_master" not in st.session_state:
    latest_df, latest_grade_master = load_latest_snapshot()
    st.session_state.df = latest_df if latest_df is not None and not latest_df.empty else generate_default()
    st.session_state.grade_master = latest_grade_master if latest_grade_master is not None and not latest_grade_master.empty else get_default_grade_master()
if "last_change_log" not in st.session_state:
    st.session_state.last_change_log = pd.DataFrame()
if "last_grade_change_log" not in st.session_state:
    st.session_state.last_grade_change_log = pd.DataFrame()

st.session_state.df = ensure_main_columns(st.session_state.df)
st.session_state.grade_master = ensure_grade_master(st.session_state.grade_master)
master_df = ensure_main_columns(st.session_state.df.copy())
grade_master_df = ensure_grade_master(st.session_state.grade_master.copy())
display_master_df = attach_grade_overview(master_df, grade_master_df)
summary_df = department_summary(master_df)

m1, m2, m3, m4 = st.columns(4)
m1.metric("部門数", len(DEPARTMENTS))
m2.metric("グレード数", len(GRADE_ORDER))
m3.metric("管理行数", len(master_df))
m4.metric("ログイン権限", "編集可" if is_admin else "閲覧のみ")

st.markdown("## フィルター")
c1, c2, c3 = st.columns([1.2, 1, 1.2])
with c1:
    dept_filter = st.multiselect("部門", options=DEPARTMENTS, default=[])
with c2:
    grade_filter = st.multiselect("グレード", options=GRADE_ORDER, default=[])
with c3:
    keyword = st.text_input("キーワード検索", placeholder="職務概要・KPI・KRA・スキル など")

filtered_df = display_master_df.copy()
if dept_filter:
    filtered_df = filtered_df[filtered_df["department"].isin(dept_filter)]
if grade_filter:
    filtered_df = filtered_df[filtered_df["grade"].isin(grade_filter)]
if keyword:
    kw = keyword.lower()
    filtered_df = filtered_df[filtered_df.apply(lambda row: any(kw in str(v).lower() for v in row.values), axis=1)]
filtered_df["grade_sort"] = filtered_df["grade"].map({g: i for i, g in enumerate(GRADE_ORDER)})
filtered_df = filtered_df.sort_values(["department", "grade_sort"]).drop(columns=["grade_sort"])
st.caption(f"表示件数: {len(filtered_df)} / 全 {len(master_df)} 件")

tab1, tab2, tab3, tab4, tab5, tab6, tab7 = st.tabs(["一覧", "編集", "配布用出力", "CSV / Excel", "履歴", "グレード定義", "設定"])

with tab1:
    st.subheader("一覧ビュー")
    view_mode = st.radio("表示形式", ["部門ごとの展開表示", "全件テーブル"], horizontal=True)
    if view_mode == "部門ごとの展開表示":
        for dept in DEPARTMENTS:
            dept_df = filtered_df[filtered_df["department"] == dept].copy()
            if dept_df.empty:
                continue
            with st.expander(f"{dept}", expanded=False):
                st.markdown(f"<div class='small-note'>{DEPARTMENT_DESCRIPTIONS.get(dept, '')}</div>", unsafe_allow_html=True)
                for _, row in dept_df.iterrows():
                    st.markdown(f"### {row['grade']} / {row['grade_name']}")
                    st.markdown(f"<div class='handout-box'>{make_handout_text(row)}</div>", unsafe_allow_html=True)
                    st.markdown("---")
    else:
        table_df = filtered_df.copy()
        table_df["配布用プレビュー"] = table_df.apply(make_handout_text, axis=1)
        st.dataframe(table_df, use_container_width=True, hide_index=True)
    st.markdown("### 部門別サマリー")
    st.dataframe(summary_df, use_container_width=True, hide_index=True)

with tab2:
    st.subheader("直接編集")
    if not is_admin:
        st.info("ユーザー権限では編集できません。閲覧とダウンロードのみ可能です。")
        st.dataframe(filtered_df, use_container_width=True, hide_index=True)
    else:
        edited_df = st.data_editor(
            filtered_df,
            use_container_width=True,
            hide_index=True,
            num_rows="fixed",
            column_config={
                "department": st.column_config.TextColumn("所属部門", disabled=True),
                "grade": st.column_config.TextColumn("グレード", disabled=True),
                "grade_name": st.column_config.TextColumn("グレード名"),
                "role_summary": st.column_config.TextColumn("職務概要", width="large"),
                "grade_overview": st.column_config.TextColumn("グレード定義", disabled=True, width="large"),
                "kpi": st.column_config.TextColumn("評価指標（KPI）", width="medium"),
                "reports_to": st.column_config.TextColumn("Reports to", width="medium"),
                "direct_reports": st.column_config.TextColumn("Direct Reports", width="medium"),
                "skills": st.column_config.TextColumn("Skills", width="large"),
                "kra": st.column_config.TextColumn("KRA", width="large"),
                "accountabilities": st.column_config.TextColumn("Key Operational Accountabilities", width="large"),
            },
        )
        memo = st.text_input("保存メモ", placeholder="例: 部門長レビュー反映")
        if st.button("編集内容を保存", type="primary"):
            new_df = update_master_from_editor(st.session_state.df, edited_df)
            st.session_state.df = ensure_main_columns(new_df)
            save_snapshot(st.session_state.df, st.session_state.grade_master, memo or "editor save")
            st.success("保存しました。")
            st.rerun()

with tab3:
    st.subheader("配布用出力")
    export_df = dataframe_for_handout_export(filtered_df)
    c1, c2, c3 = st.columns(3)
    with c1:
        st.download_button("配布用Excel", data=create_handout_excel_bytes(export_df), file_name="job_matrix_handout.xlsx", mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")
    with c2:
        st.download_button("配布用Word", data=create_handout_docx_bytes(export_df), file_name="job_matrix_handout.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
    with c3:
        st.download_button("配布用PDF", data=create_handout_pdf_bytes(export_df), file_name="job_matrix_handout.pdf", mime="application/pdf")

with tab4:
    st.subheader("CSV / Excel")
    d1, d2, d3, d4, d5 = st.columns(5)
    with d1:
        st.download_button("現在データCSV", data=dataframe_to_csv_bytes(st.session_state.df), file_name="job_matrix_export.csv", mime="text/csv")
    with d2:
        st.download_button("現在データExcel", data=dataframe_to_excel_bytes(st.session_state.df, st.session_state.grade_master), file_name="job_matrix_export.xlsx", mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")
    with d3:
        st.download_button("グレード定義CSV", data=grade_master_to_csv_bytes(st.session_state.grade_master), file_name="grade_master_export.csv", mime="text/csv")
    with d4:
        st.download_button("業務マスタ空テンプレートCSV", data=minimal_template_csv_bytes(), file_name="job_matrix_template.csv", mime="text/csv")
    with d5:
        st.download_button("グレード定義テンプレートCSV", data=minimal_grade_template_csv_bytes(), file_name="grade_master_template.csv", mime="text/csv")

    if not is_admin:
        st.info("ユーザー権限ではCSV取込はできません。ダウンロードのみ可能です。")
    else:
        st.markdown("---")
        st.markdown("### 業務マスタCSV部分更新")
        uploaded_csv = st.file_uploader("業務マスタ更新用CSV", type=["csv"], key="job_csv")
        if uploaded_csv is not None:
            try:
                import_df = pd.read_csv(uploaded_csv)
                st.dataframe(import_df, use_container_width=True, hide_index=True)
                if "department" in import_df.columns and "grade" in import_df.columns:
                    invalid_df = validate_import_keys(st.session_state.df, import_df)
                    if not invalid_df.empty:
                        st.warning("マスタに存在しない department + grade が含まれています。")
                        st.dataframe(invalid_df, use_container_width=True, hide_index=True)
                    else:
                        st.success("すべてのキーが既存マスタに存在しています。")
                csv_memo = st.text_input("業務マスタ更新メモ", key="job_csv_memo")
                if st.button("業務マスタをCSV内容で部分更新する", type="primary"):
                    new_df, change_log = merge_partial_update(st.session_state.df, import_df)
                    st.session_state.df = ensure_main_columns(new_df)
                    st.session_state.last_change_log = change_log
                    save_snapshot(st.session_state.df, st.session_state.grade_master, csv_memo or f"job csv partial update ({len(change_log)} changes)")
                    st.success(f"更新完了: {len(change_log)} 件")
                    st.rerun()
            except Exception as e:
                st.error(f"CSV取込エラー: {e}")

        st.markdown("---")
        st.markdown("### グレード定義CSV部分更新")
        uploaded_grade_csv = st.file_uploader("グレード定義更新用CSV", type=["csv"], key="grade_csv")
        if uploaded_grade_csv is not None:
            try:
                import_grade_df = pd.read_csv(uploaded_grade_csv)
                st.dataframe(import_grade_df, use_container_width=True, hide_index=True)
                grade_csv_memo = st.text_input("グレード定義更新メモ", key="grade_csv_memo")
                if st.button("グレード定義をCSV内容で部分更新する", type="primary"):
                    new_grade_df, grade_change_log = merge_grade_master_update(st.session_state.grade_master, import_grade_df)
                    st.session_state.grade_master = ensure_grade_master(new_grade_df)
                    st.session_state.last_grade_change_log = grade_change_log
                    save_snapshot(st.session_state.df, st.session_state.grade_master, grade_csv_memo or f"grade csv partial update ({len(grade_change_log)} changes)")
                    st.success(f"更新完了: {len(grade_change_log)} 件")
                    st.rerun()
            except Exception as e:
                st.error(f"グレード定義CSV取込エラー: {e}")

with tab5:
    st.subheader("履歴")
    history = load_history()
    hist_df = pd.DataFrame(history)
    if hist_df.empty:
        st.info("履歴はまだありません。")
    else:
        show_cols = [c for c in ["id", "memo", "created_at"] if c in hist_df.columns]
        st.dataframe(hist_df[show_cols], use_container_width=True, hide_index=True)
        if is_admin and "id" in hist_df.columns:
            options = [f"{r['id']} | {r.get('memo', '')} | {r.get('created_at', '')}" for _, r in hist_df.iterrows()]
            selected = st.selectbox("復元対象", options=options)
            if st.button("この履歴を復元する"):
                selected_id = selected.split(" | ")[0]
                target = next((x for x in history if str(x.get("id")) == selected_id), None)
                if target:
                    st.session_state.df = ensure_main_columns(pd.DataFrame(target.get("data", [])))
                    st.session_state.grade_master = ensure_grade_master(pd.DataFrame(target.get("grade_master", [])))
                    save_snapshot(st.session_state.df, st.session_state.grade_master, f"restore from {selected_id}")
                    st.success("復元しました。")
                    st.rerun()
        elif not is_admin:
            st.info("ユーザー権限では履歴復元はできません。")

with tab6:
    st.subheader("グレード定義")
    if is_admin:
        edited_grade_df = st.data_editor(
            st.session_state.grade_master,
            use_container_width=True,
            hide_index=True,
            num_rows="fixed",
            column_config={
                "grade": st.column_config.TextColumn("グレード", disabled=True),
                "grade_name": st.column_config.TextColumn("グレード名"),
                "grade_overview": st.column_config.TextColumn("グレード定義", width="large"),
            },
        )
        grade_memo = st.text_input("グレード保存メモ", key="grade_editor_memo")
        if st.button("グレード定義を保存"):
            st.session_state.grade_master = ensure_grade_master(edited_grade_df)
            save_snapshot(st.session_state.df, st.session_state.grade_master, grade_memo or "grade editor save")
            st.success("保存しました。")
            st.rerun()
    else:
        st.dataframe(st.session_state.grade_master, use_container_width=True, hide_index=True)

with tab7:
    st.subheader("設定")
    if is_admin:
        st.markdown("### ユーザー管理")
        st.caption(f"このアプリは専用のログインテーブル `{USERS_TABLE}` を使います。別アプリと分けて運用できます。")
        users_df = pd.DataFrame(rest_get(USERS_TABLE, params={"select": "id,username,display_name,role,is_active,failed_attempts,locked_until,created_at,updated_at", "order": "created_at.asc"}))
        if not users_df.empty:
            st.dataframe(users_df, use_container_width=True, hide_index=True)
        with st.expander("新規ユーザー追加"):
            with st.form("create_user_form"):
                new_username = st.text_input(LOGIN_ID_LABEL)
                new_display_name = st.text_input("表示名")
                new_role = st.selectbox("権限", ["viewer", "admin"])
                new_password = st.text_input("初期パスワード", type="password")
                created = st.form_submit_button("作成", type="primary")
            if created:
                if not new_username or not new_password:
                    st.error(f"{LOGIN_ID_LABEL} とパスワードは必須です。")
                elif get_user_by_username(new_username):
                    st.error(f"その{LOGIN_ID_LABEL}は既に存在します。")
                else:
                    rest_post(USERS_TABLE, {
                        "username": normalize_username(new_username),
                        "display_name": new_display_name or normalize_username(new_username),
                        "role": new_role,
                        "password_hash": hash_password(new_password),
                        "is_active": True,
                        "failed_attempts": 0,
                        "locked_until": None,
                        "created_at": now_utc().isoformat(),
                        "updated_at": now_utc().isoformat(),
                    })
                    st.success("ユーザーを作成しました。")
                    st.rerun()

        with st.expander("既存ユーザー更新 / ロック解除 / パスワード再設定"):
            user_options = [u["username"] for u in rest_get(USERS_TABLE, params={"select": "username", "order": "username.asc"})]
            target_username = st.selectbox("対象ユーザー", user_options)
            target_user = get_user_by_username(target_username) if target_username else None
            if target_user:
                col1, col2 = st.columns(2)
                with col1:
                    upd_display_name = st.text_input("表示名", value=target_user.get("display_name") or "", key="upd_display_name")
                    upd_role = st.selectbox("権限", ["viewer", "admin"], index=0 if target_user.get("role") == "viewer" else 1, key="upd_role")
                    upd_active = st.checkbox("有効", value=bool(target_user.get("is_active", True)), key="upd_active")
                    if st.button("基本情報を更新"):
                        update_user_row(target_username, {"display_name": upd_display_name or target_username, "role": upd_role, "is_active": upd_active, "updated_at": now_utc().isoformat()})
                        st.success("更新しました。")
                        st.rerun()
                with col2:
                    reset_pw = st.text_input("新しいパスワード", type="password", key="reset_pw")
                    if st.button("パスワード再設定"):
                        if not reset_pw:
                            st.error("新しいパスワードを入力してください。")
                        else:
                            update_user_row(target_username, {"password_hash": hash_password(reset_pw), "updated_at": now_utc().isoformat()})
                            st.success("パスワードを更新しました。")
                            st.rerun()
                    if st.button("ロック解除"):
                        clear_failed_attempts(target_username)
                        st.success("ロック解除しました。")
                        st.rerun()
                    if target_username != auth_user.get("username") and st.button("このユーザーを削除"):
                        rest_delete(USERS_TABLE, params={"username": f"eq.{target_username}"})
                        st.success("削除しました。")
                        st.rerun()
    else:
        st.info("ユーザー権限では設定変更はできません。")
        st.write("閲覧とダウンロードのみ利用できます。")
